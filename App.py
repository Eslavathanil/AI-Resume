import streamlit as st
import pdfplumber
import docx
import spacy
import re
import phonenumbers
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load NLP model
nlp = spacy.load("en_core_web_sm")

# Predefined skill set
SKILL_SET = {
    "Python", "Java", "C++", "Machine Learning", "Data Science", "Deep Learning",
    "SQL", "Django", "Flask", "TensorFlow", "Keras", "Pandas", "NumPy",
    "Natural Language Processing", "Computer Vision", "JavaScript", "React", "Node.js"
}

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    text = ""
    page_count = 0
    with pdfplumber.open(pdf_file) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text, page_count

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    page_count = len(doc.paragraphs) // 40  # Approximate 40 paragraphs per page
    return text, max(1, page_count)

# Function to extract personal details (Name, Phone, Email)
def extract_personal_info(text):
    name = None
    email = None
    phone = None

    # Extract name (assumption: first two words are the name)
    words = text.split()
    if len(words) > 1:
        name = words[0] + " " + words[1]  

    # Extract email
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    if email_match:
        email = email_match.group(0)

    # Extract phone number
    phone_matches = re.findall(r"\+?\d[\d -]{8,15}\d", text)
    for match in phone_matches:
        parsed_phone = phonenumbers.parse(match, "IN")  # Assume Indian numbers
        if phonenumbers.is_valid_number(parsed_phone):
            phone = match
            break

    return name, email, phone

# Function to extract skills
def extract_skills(text):
    doc = nlp(text.lower())
    found_skills = {token.text for token in doc if token.text in SKILL_SET}
    return list(found_skills)

# Function to calculate resume-job description match score
def calculate_match_score(resume_text, job_desc_text):
    vectorizer = CountVectorizer().fit_transform([resume_text, job_desc_text])
    similarity = cosine_similarity(vectorizer[0], vectorizer[1])
    return round(similarity[0][0] * 100, 2)

# Apply Custom Styling
st.markdown("""
    <style>
        body {
            background-color: #121212;
            color: white;
        }
        .css-1d391kg {
            background-color: #1E1E1E !important;
        }
        .stTextInput>div>div>input {
            color: white !important;
            background-color: #333333 !important;
            border: 1px solid #555555 !important;
        }
        .stTextArea>div>textarea {
            color: white !important;
            background-color: #333333 !important;
            border: 1px solid #555555 !important;
        }
        .stButton>button {
            color: white !important;
            background-color: #6200EA !important;
            border: none !important;
        }
    </style>
""", unsafe_allow_html=True)

# Sidebar with App Maker Name
st.sidebar.title("🔹 AI Resume Screening")
st.sidebar.markdown("👨‍💻 **Developed by : ESLAVATH ANIL**")  # Your Name

# Main App Title
st.title(" AI-Powered Resume Screening System")
st.markdown("**An intelligent system to evaluate resumes and match them with job descriptions.**")

# File uploader and job description input
uploaded_resume = st.file_uploader("📂 Upload your Resume (PDF or DOCX)", type=["pdf", "docx"])
job_desc = st.text_area("📝 Paste the Job Description", height=150)

# Select Fresher or Experienced
experience_level = st.radio("🎯 Select Experience Level", ("Fresher", "Experienced"))

# Analyze button
if st.button("🔍 Analyze Resume") and uploaded_resume and job_desc:
    st.subheader(" Resume Analysis Results")

    # Extract text & page count
    if uploaded_resume.type == "application/pdf":
        resume_text, resume_pages = extract_text_from_pdf(uploaded_resume)
    elif uploaded_resume.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume_text, resume_pages = extract_text_from_docx(uploaded_resume)

    # Extract personal information
    name, email, phone = extract_personal_info(resume_text)

    # Extract skills from resume
    resume_skills = extract_skills(resume_text)

    # Extract skills from job description
    job_skills = extract_skills(job_desc)

    # Calculate match score
    match_score = calculate_match_score(resume_text, job_desc)

    # Determine missing skills
    missing_skills = list(set(job_skills) - set(resume_skills))

    # Display Results
    col1, col2 = st.columns([2, 2])

    with col1:
        st.markdown("### 📝 Basic Information")
        st.write(f"📃 **Resume Pages:** {resume_pages}")
        st.write(f"👤 **Name:** {name if name else 'Not Found'}")
        st.write(f"📧 **Email:** {email if email else 'Not Found'}")
        st.write(f"📞 **Phone:** {phone if phone else 'Not Found'}")

    with col2:
        st.markdown("### 🔍 Resume Analysis")
        st.write(f"✅ **Match Score:** {match_score}%")
        
        if resume_skills:
            st.write(f"🛠 **Skills Found in Resume:** {', '.join(resume_skills)}")
        if job_skills:
            st.write(f"📌 **Required Job Skills:** {', '.join(job_skills)}")

    # Missing Skills Section
    if missing_skills:
        st.markdown("### ⚠️ Missing Skills & Recommendations")
        st.warning(f"Your resume is missing the following required skills: {', '.join(missing_skills)}")
        st.markdown("💡 **How to Improve:**")
        for skill in missing_skills:
            st.markdown(f"🔹 Gain expertise in **{skill}** through online courses or projects.")
    else:
        st.success("🎉 Your resume fully matches the job requirements! Great job!")